from docx import Document
from docx.shared import Pt

from claas.curriculum_converter import CurriculumConverter


class WordConverter(CurriculumConverter):
    def start_output(self, title):
        document = Document()
        document.add_heading(title, level=1)
        return document

    def start_module(self, output, title: str, description: str, total_time: int):
        if self.include_time and total_time:
            output.add_heading(f"{title} ({total_time} UE)", level=2)
        else:
            output.add_heading(title, level=2)
        if description:
            output.add_paragraph(description)

    def add_summary_topic(
        self, output, contents: str, duration: str, method: str, material: str
    ):
        paragraph = output.add_paragraph(style="List Bullet")
        if self.include_time and duration:
            paragraph.text = f"{contents} ({duration} UE)"
        else:
            paragraph.text = contents

    def add_detail_topic(
        self, output, contents: str, duration: str, method: str, material: str
    ):
        paragraph = output.add_paragraph(style="List Bullet 2")
        if self.include_time and duration:
            paragraph.text = f"{contents} ({duration} UE)"
        else:
            paragraph.text = contents
        paragraph.style.font.size = Pt(10)  # Make subtopics slightly smaller

    def add_section(self, output, text: str, week_time: int = None):
        if self.include_time and week_time:
            output.add_heading(f"{text} ({week_time} UE)", level=3)
        else:
            output.add_heading(text, level=3)

    def finalize_output(self, output):
        if self.include_time:
            p = output.add_paragraph()
            run = p.add_run(
                f"Unterrichtseinheiten insgesamt: {self.total_course_hours}"
            )
            run.bold = True
        return output

    def save(self, output_path):
        doc = self.convert()
        doc.save(output_path)
