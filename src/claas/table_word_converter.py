import warnings

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Inches

from claas.curriculum_converter import CurriculumConverter


class TableWordConverter(CurriculumConverter):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._current_table = None
        self._current_table_total_duration = 0

    def start_output(self, title):
        document = Document()

        # Set page size to A4
        section = document.sections[0]
        section.page_height = Cm(29.7)  # A4 height in cm
        section.page_width = Cm(21.0)  # A4 width in cm

        # Set margins
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

        document.add_heading(title, level=1)
        return document

    def start_module(self, output, title: str, description: str, total_time: int):
        if self._current_table:
            self._add_table_footer(output)
            output.add_page_break()

        if self.include_time and total_time:
            output.add_heading(f"{title} ({total_time} UE)", level=2)
        else:
            output.add_heading(title, level=2)
        if description:
            output.add_paragraph(description)

    def add_topics(self, output, topics):
        if topics:
            self._current_table = output.add_table(rows=1, cols=4)
            self._current_table.style = "Table Grid"
            self._current_table.autofit = False
            self._current_table.allow_autofit = False

            self._set_table_width(
                self._current_table, Cm(12)
            )  # Set table width to 17 cm

            hdr_cells = self._current_table.rows[0].cells
            hdr_cells[0].text = "Inhalt"
            hdr_cells[1].text = "UE"
            hdr_cells[2].text = "Methodik/Didaktik"
            hdr_cells[3].text = "Materialien"

            # Set column widths
            self._current_table.columns[0].width = Cm(7)
            self._current_table.columns[1].width = Cm(1)
            self._current_table.columns[2].width = Cm(4)
            self._current_table.columns[3].width = Cm(4)

            self._set_header_style(hdr_cells)
            self._set_cell_margins(hdr_cells)

            self._current_table_total_duration = 0

            for topic_type, topic in topics:
                if topic_type == "summary":
                    self.add_summary_topic(self._current_table, *topic)
                else:
                    self.add_detail_topic(self._current_table, *topic)

            self._add_table_footer(output)

    @staticmethod
    def _set_table_width(table, width):
        table.width = width
        for cell in table.rows[0].cells:
            cell.width = width / len(table.columns)

    def add_summary_topic(
        self, table, contents: str, duration: str, method: str, material: str
    ):
        self._add_topic(table, contents, duration, method, material, is_summary=True)

    def add_detail_topic(
        self, table, contents: str, duration: str, method: str, material: str
    ):
        self._add_topic(table, contents, duration, method, material, is_summary=False)

    def _add_topic(
        self,
        table,
        contents: str,
        duration: str,
        method: str,
        material: str,
        is_summary: bool,
    ):
        try:
            duration_value = int(duration)
        except ValueError:
            duration_value = 0
            duration += " (ungültig)"
        self._current_table_total_duration += duration_value
        row_cells = table.add_row().cells
        row_cells[0].text = contents
        if is_summary:
            row_cells[0].paragraphs[0].runs[0].bold = True
        else:
            row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
            row_cells[0].paragraphs[0].style.paragraph_format.left_indent = Pt(10)
        row_cells[1].text = duration
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[2].text = method
        row_cells[3].text = material
        self._set_cell_margins(row_cells)

    def add_section(self, output, text: str, week_time: int = None):
        if self._current_table:
            self._add_table_footer(output)

        # Add extra space before the week title
        spacer = output.add_paragraph()
        spacer.space_after = Pt(6)  # Adjust this value to increase or decrease space

        p = output.add_paragraph()
        if self.include_time and week_time:
            text = f"{text} ({week_time} UE)"
        run = p.add_run(text)
        run.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Optionally, add some space after the week title as well
        p.space_after = Pt(3)  # Adjust this value as needed

    def finalize_output(self, output):
        if self._current_table:
            self._add_table_footer(output)
        if self.include_time:
            p = output.add_paragraph()
            run = p.add_run(
                f"Unterrichtseinheiten insgesamt: {self.total_course_hours}"
            )
            run.bold = True
        return output

    def save(self, output_path):
        doc = self.convert()
        doc.save(output_path)

    def _add_table_footer(self, output):
        self._current_table.add_row()
        row_cells = self._current_table.rows[-1].cells
        row_cells[0].text = "Summe:"
        row_cells[1].text = str(self._current_table_total_duration)
        for i in range(2, 4):
            row_cells[i].text = ""
        self._set_footer_style(row_cells)
        row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        self._set_cell_margins(row_cells)
        self._current_table = None

    @staticmethod
    def _set_header_style(cells):
        assert cells is not None, "Cells must not be None to set header style."
        for cell in cells:
            cell_paragraph = cell.paragraphs[0]
            cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = cell_paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(12)
            cell_shading = OxmlElement("w:shd")
            cell_shading.set(qn("w:fill"), "D9D9D9")  # Light grey background
            cell._element.get_or_add_tcPr().append(cell_shading)

    @staticmethod
    def _set_footer_style(cells):
        assert cells is not None, "Cells must not be None to set footer style."
        for cell in cells:
            cell_paragraph = cell.paragraphs[0]
            cell_paragraph.alignment = (
                WD_PARAGRAPH_ALIGNMENT.RIGHT
                if cell.text == "Summe:"
                else WD_PARAGRAPH_ALIGNMENT.CENTER
            )
            run = cell_paragraph.runs[0]
            run.bold = False
            run.font.size = Pt(11)
            cell_shading = OxmlElement("w:shd")
            cell_shading.set(qn("w:fill"), "D9D9D9")  # Light grey background
            cell._element.get_or_add_tcPr().append(cell_shading)

    @staticmethod
    def _set_cell_margins(cells, top=100, start=100, bottom=100, end=100):
        assert cells is not None, "Cells must not be None to set cell margins."
        for cell in cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement("w:tcMar")
            for margin, value in [
                ("top", top),
                ("start", start),
                ("bottom", bottom),
                ("end", end),
            ]:
                node = OxmlElement(f"w:{margin}")
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")
                tcMar.append(node)
            tcPr.append(tcMar)
